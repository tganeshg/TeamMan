from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from datetime import date
import io

from database import get_db
from models import Release, Task, TeamMember
from schemas import ReleaseReport, MemberReportRow, ReportTaskRow

router = APIRouter(prefix="/reports", tags=["reports"])

_TERMINAL = {"SID12", "SID13"}
_EARLY_DAYS = 2  # closed more than 2 days before end_date = "early"

STATUS_LABEL = {
    "SID00": "Not Started", "SID01": "Study", "SID02": "Requirement",
    "SID03": "POC", "SID04": "Core Impl", "SID05": "Dev Testing",
    "SID06": "Review", "SID07": "Rework", "SID08": "Ready to Merge",
    "SID09": "Ready to Release", "SID10": "Waiting", "SID11": "Reopened",
    "SID12": "Closed", "SID13": "Released", "SID14": "On Hold",
    "SID15": "Debug", "SID16": "Moved to Software",
}


def _task_timing(task: Task) -> str:
    """Return timing label for a task."""
    today = date.today()
    closed = task.status in _TERMINAL
    if not task.end_date:
        return "closed" if closed else "open"
    if closed and task.closed_at:
        diff = (task.end_date - task.closed_at).days
        if diff > _EARLY_DAYS:
            return "early"
        elif diff >= 0:
            return "on_time"
        else:
            return "overdue"
    elif closed:
        return "on_time"  # closed but no closed_at recorded
    else:
        if task.end_date < today:
            return "overdue_open"
        return "in_progress"


def _build_report(release: Release, db: Session) -> ReleaseReport:
    tasks = db.query(Task).filter(Task.release_id == release.id).all()
    members = db.query(TeamMember).all()
    today = date.today()

    # Per-member stats
    member_map: dict[int, MemberReportRow] = {}
    for m in members:
        member_map[m.id] = MemberReportRow(
            member_id=m.id, member_name=m.name, member_role=m.role,
            total_assigned=0, closed=0, on_time=0, early=0,
            overdue_closed=0, still_open=0, in_progress=0,
        )

    task_rows: list[ReportTaskRow] = []
    totals = dict(total=0, closed=0, on_time=0, early=0,
                  overdue_closed=0, still_open=0, in_progress=0)

    for t in tasks:
        timing = _task_timing(t)
        is_closed = t.status in _TERMINAL
        assignee_name = t.assignee.name if t.assignee else None
        totals["total"] += 1

        if is_closed:
            totals["closed"] += 1
            if timing == "early":
                totals["early"] += 1
            elif timing == "on_time":
                totals["on_time"] += 1
            elif timing == "overdue":
                totals["overdue_closed"] += 1
        else:
            if timing == "overdue_open":
                totals["still_open"] += 1
            else:
                totals["in_progress"] += 1

        if t.assignee_id and t.assignee_id in member_map:
            row = member_map[t.assignee_id]
            row.total_assigned += 1
            if is_closed:
                row.closed += 1
                if timing == "early":
                    row.early += 1
                elif timing == "on_time":
                    row.on_time += 1
                elif timing == "overdue":
                    row.overdue_closed += 1
            else:
                if timing == "overdue_open":
                    row.still_open += 1
                else:
                    row.in_progress += 1

        display_timing = timing if timing != "overdue_open" else "overdue"
        task_rows.append(ReportTaskRow(
            task_id=t.id,
            portal_task_id=t.portal_task_id,
            title=t.title,
            task_type=t.task_type,
            assignee_name=assignee_name,
            status=f"{t.status} – {STATUS_LABEL.get(t.status, t.status)}",
            end_date=t.end_date,
            closed_at=t.closed_at,
            timing=display_timing,
        ))

    # Only include members with assigned tasks
    member_rows = [r for r in member_map.values() if r.total_assigned > 0]
    member_rows.sort(key=lambda r: r.member_name)

    return ReleaseReport(
        release_id=release.id,
        release_name=release.name,
        release_date=release.release_date,
        release_status=release.status,
        generated_on=today,
        total_tasks=totals["total"],
        total_closed=totals["closed"],
        total_on_time=totals["on_time"],
        total_early=totals["early"],
        total_overdue_closed=totals["overdue_closed"],
        total_still_open=totals["still_open"],
        total_in_progress=totals["in_progress"],
        members=member_rows,
        tasks=task_rows,
    )


@router.get("/{release_id}", response_model=ReleaseReport)
def get_report(release_id: int, db: Session = Depends(get_db)):
    release = db.query(Release).filter(Release.id == release_id).first()
    if not release:
        raise HTTPException(status_code=404, detail="Release not found")
    return _build_report(release, db)


@router.get("/{release_id}/export/pdf")
def export_pdf(release_id: int, db: Session = Depends(get_db)):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    release = db.query(Release).filter(Release.id == release_id).first()
    if not release:
        raise HTTPException(status_code=404, detail="Release not found")

    report = _build_report(release, db)
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    # ── Colours
    PRIMARY   = colors.HexColor("#4e73df")
    SUCCESS   = colors.HexColor("#1cc88a")
    DANGER    = colors.HexColor("#e74a3b")
    WARNING   = colors.HexColor("#f6c23e")
    MUTED     = colors.HexColor("#858796")
    DARK      = colors.HexColor("#3a3f5c")
    LIGHT_BG  = colors.HexColor("#f8f9fc")
    HDR_BG    = colors.HexColor("#4e73df")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", fontSize=18, fontName="Helvetica-Bold",
                                  textColor=PRIMARY, alignment=TA_CENTER, spaceAfter=4)
    sub_style   = ParagraphStyle("sub",   fontSize=9,  fontName="Helvetica",
                                  textColor=MUTED, alignment=TA_CENTER, spaceAfter=2)
    section_style = ParagraphStyle("sec", fontSize=11, fontName="Helvetica-Bold",
                                    textColor=DARK, spaceBefore=12, spaceAfter=4)
    cell_style  = ParagraphStyle("cell",  fontSize=8,  fontName="Helvetica", leading=10)
    footer_style = ParagraphStyle("footer", fontSize=7, fontName="Helvetica",
                                   textColor=MUTED, alignment=TA_CENTER, spaceBefore=12)

    def hdr_cell(text):
        return Paragraph(f"<b>{text}</b>", ParagraphStyle(
            "hc", fontSize=8, fontName="Helvetica-Bold",
            textColor=colors.white, alignment=TA_CENTER, leading=10))

    def val_cell(text, align=TA_CENTER, color=DARK):
        return Paragraph(str(text), ParagraphStyle(
            "vc", fontSize=8, fontName="Helvetica",
            textColor=color, alignment=align, leading=10))

    timing_label = {
        "early": "Early ✓", "on_time": "On Time ✓",
        "overdue": "Overdue ✗", "in_progress": "In Progress",
        "open": "Open", "closed": "Closed",
    }
    timing_color = {
        "early": SUCCESS, "on_time": PRIMARY,
        "overdue": DANGER, "in_progress": WARNING,
        "open": MUTED, "closed": SUCCESS,
    }

    elements = []

    # ── Title
    elements.append(Paragraph("Prime Team Report", title_style))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(Paragraph(
        f"Release: <b>{report.release_name}</b>  |  "
        f"Target: <b>{report.release_date.strftime('%d %b %Y')}</b>  |  "
        f"Generated: <b>{report.generated_on.strftime('%d %b %Y')}</b>",
        sub_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=PRIMARY, spaceAfter=8))

    # ── Summary
    elements.append(Paragraph("Summary", section_style))
    summary_rows = [[
        hdr_cell("Total Tasks"), hdr_cell("Closed"), hdr_cell("Early"),
        hdr_cell("On Time"), hdr_cell("Late Closed"), hdr_cell("Open / Overdue"), hdr_cell("In Progress"),
    ], [
        val_cell(report.total_tasks),
        val_cell(report.total_closed),
        val_cell(report.total_early),
        val_cell(report.total_on_time),
        val_cell(report.total_overdue_closed),
        val_cell(report.total_still_open),
        val_cell(report.total_in_progress),
    ]]
    summary_tbl = Table(summary_rows, colWidths=[2.4*cm]*7)
    summary_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), HDR_BG),
        ("BACKGROUND", (0,1), (-1,1), LIGHT_BG),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#e3e6f0")),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [LIGHT_BG, colors.white]),
        ("TOPPADDING",  (0,0), (-1,-1), 5),
        ("BOTTOMPADDING", (0,0), (-1,-1), 5),
    ]))
    elements.append(summary_tbl)

    # ── Per-member breakdown
    elements.append(Paragraph("Per-Member Breakdown", section_style))
    member_header = [
        hdr_cell("Member"), hdr_cell("Role"), hdr_cell("Assigned"),
        hdr_cell("Closed"), hdr_cell("Early"), hdr_cell("On Time"),
        hdr_cell("Late Closed"), hdr_cell("Open/OD"), hdr_cell("In Progress"),
    ]
    member_data = [member_header]
    for m in report.members:
        member_data.append([
            val_cell(m.member_name, TA_LEFT),
            val_cell(m.member_role),
            val_cell(m.total_assigned),
            val_cell(m.closed),
            val_cell(m.early,         color=SUCCESS),
            val_cell(m.on_time,       color=PRIMARY),
            val_cell(m.overdue_closed, color=DANGER),
            val_cell(m.still_open,    color=DANGER),
            val_cell(m.in_progress,   color=WARNING),
        ])
    col_w = [3.5*cm, 2*cm, 1.8*cm, 1.8*cm, 1.6*cm, 1.6*cm, 2*cm, 1.8*cm, 2*cm]
    member_tbl = Table(member_data, colWidths=col_w)
    member_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0), HDR_BG),
        ("GRID",          (0,0), (-1,-1), 0.5, colors.HexColor("#e3e6f0")),
        ("ALIGN",         (0,0), (-1,-1), "CENTER"),
        ("ALIGN",         (0,1), (0,-1), "LEFT"),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [LIGHT_BG, colors.white]),
        ("TOPPADDING",    (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ("FONTNAME",      (0,1), (0,-1), "Helvetica-Bold"),
    ]))
    elements.append(member_tbl)

    # ── Task details
    elements.append(Paragraph("Task Details", section_style))
    task_header = [
        hdr_cell("ID"), hdr_cell("Title"), hdr_cell("Type"),
        hdr_cell("Assignee"), hdr_cell("Status"),
        hdr_cell("Due Date"), hdr_cell("Closed On"), hdr_cell("Timing"),
    ]
    task_data = [task_header]
    for t in report.tasks:
        tc = timing_color.get(t.timing, MUTED)
        task_data.append([
            val_cell(f"#{t.portal_task_id}" if t.portal_task_id else f"#{t.task_id}"),
            val_cell(t.title, TA_LEFT),
            val_cell(t.task_type.capitalize()),
            val_cell(t.assignee_name or "—"),
            val_cell(t.status, TA_LEFT),
            val_cell(t.end_date.strftime("%d %b %Y") if t.end_date else "—"),
            val_cell(t.closed_at.strftime("%d %b %Y") if t.closed_at else "—"),
            val_cell(timing_label.get(t.timing, t.timing), color=tc),
        ])
    task_col_w = [1.5*cm, 5.5*cm, 1.8*cm, 2.5*cm, 3*cm, 2.2*cm, 2.2*cm, 2*cm]
    task_tbl = Table(task_data, colWidths=task_col_w, repeatRows=1)
    task_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0), HDR_BG),
        ("GRID",          (0,0), (-1,-1), 0.5, colors.HexColor("#e3e6f0")),
        ("ALIGN",         (0,0), (-1,-1), "CENTER"),
        ("ALIGN",         (1,1), (1,-1), "LEFT"),
        ("ALIGN",         (4,1), (4,-1), "LEFT"),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [LIGHT_BG, colors.white]),
        ("TOPPADDING",    (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    elements.append(task_tbl)

    # ── Footer
    elements.append(Paragraph(
        f"Report generated by PrimeDesk  |  Prime Team  |  {report.generated_on.strftime('%d %b %Y')}",
        footer_style))

    doc.build(elements)
    buf.seek(0)

    filename = f"PrimeDesk_Report_{release.name.replace(' ', '_')}_{report.generated_on}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{release_id}/export/docx")
def export_docx(release_id: int, db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    release = db.query(Release).filter(Release.id == release_id).first()
    if not release:
        raise HTTPException(status_code=404, detail="Release not found")

    report = _build_report(release, db)
    doc = Document()

    # ── Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Title
    title = doc.add_heading("Prime Team Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x4E, 0x73, 0xDF)

    sub = doc.add_paragraph(f"Release: {report.release_name}  |  Target Date: {report.release_date.strftime('%d %b %Y')}  |  Generated: {report.generated_on.strftime('%d %b %Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x5A, 0x5C, 0x69)

    doc.add_paragraph()

    # ── Summary section
    doc.add_heading("Summary", level=1)
    summary_data = [
        ("Total Tasks", str(report.total_tasks)),
        ("Closed", str(report.total_closed)),
        ("Closed Early", str(report.total_early)),
        ("Closed On Time", str(report.total_on_time)),
        ("Closed Overdue", str(report.total_overdue_closed)),
        ("Still Open (overdue)", str(report.total_still_open)),
        ("In Progress", str(report.total_in_progress)),
    ]
    tbl = doc.add_table(rows=1, cols=len(summary_data))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, (label, val) in enumerate(summary_data):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}\n{val}")
        run.font.size = Pt(9)
        run.font.bold = True

    doc.add_paragraph()

    # ── Per-member breakdown
    doc.add_heading("Per-Member Breakdown", level=1)
    cols = ["Member", "Role", "Assigned", "Closed", "Early", "On Time", "Late Closed", "Open/Overdue", "In Progress"]
    tbl2 = doc.add_table(rows=1 + len(report.members), cols=len(cols))
    tbl2.style = "Table Grid"
    for i, col in enumerate(cols):
        cell = tbl2.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(col)
        run.font.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, m in enumerate(report.members):
        cells = tbl2.rows[row_idx + 1].cells
        values = [m.member_name, m.member_role, str(m.total_assigned),
                  str(m.closed), str(m.early), str(m.on_time),
                  str(m.overdue_closed), str(m.still_open), str(m.in_progress)]
        for col_idx, val in enumerate(values):
            p = cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size = Pt(9)

    doc.add_paragraph()

    # ── Task detail table
    doc.add_heading("Task Details", level=1)
    task_cols = ["ID", "Portal ID", "Title", "Type", "Assignee", "Status", "Due Date", "Closed On", "Timing"]
    tbl3 = doc.add_table(rows=1 + len(report.tasks), cols=len(task_cols))
    tbl3.style = "Table Grid"
    for i, col in enumerate(task_cols):
        cell = tbl3.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(col)
        run.font.bold = True
        run.font.size = Pt(8)

    timing_display = {
        "early": "Early ✓", "on_time": "On Time ✓",
        "overdue": "Overdue ✗", "in_progress": "In Progress",
        "open": "Open", "closed": "Closed",
    }

    for row_idx, t in enumerate(report.tasks):
        cells = tbl3.rows[row_idx + 1].cells
        values = [
            str(t.task_id),
            f"#{t.portal_task_id}" if t.portal_task_id else "—",
            t.title,
            t.task_type.capitalize(),
            t.assignee_name or "—",
            t.status,
            t.end_date.strftime("%d %b %Y") if t.end_date else "—",
            t.closed_at.strftime("%d %b %Y") if t.closed_at else "—",
            timing_display.get(t.timing, t.timing),
        ]
        for col_idx, val in enumerate(values):
            p = cells[col_idx].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)

    # ── Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"Report generated by PrimeDesk  |  Prime Team  |  {report.generated_on.strftime('%d %b %Y')}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x9E, 0x9F, 0xB4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"PrimeDesk_Report_{release.name.replace(' ', '_')}_{report.generated_on}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
